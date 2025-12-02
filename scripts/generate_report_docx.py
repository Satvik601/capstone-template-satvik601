# scripts/generate_report_docx.py
import json
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def write_consulting_report(final_report: dict, out_path: str):
    doc = Document()
    
    # ========== TITLE PAGE ==========
    title = doc.add_heading(level=0)
    title_run = title.add_run("Business Consulting Report")
    title_run.font.size = Pt(24)
    title_run.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(final_report["business_snapshot"].get("description", "").title())
    subtitle_run.font.size = Pt(16)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()  # spacing
    generated = doc.add_paragraph()
    generated_run = generated.add_run(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    generated_run.font.size = Pt(10)
    generated_run.italic = True
    generated.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_page_break()
    
    # ========== EXECUTIVE SUMMARY ==========
    doc.add_heading("Executive Summary", level=1)
    
    doc.add_heading("Business Overview", level=2)
    doc.add_paragraph(f"Business: {final_report['business_snapshot'].get('description', '')}")
    doc.add_paragraph(f"Primary Goal: {final_report['business_snapshot'].get('goal', '')}")
    
    doc.add_heading("Current KPIs", level=2)
    kpis = final_report["business_snapshot"].get("kpis", {})
    if kpis:
        for k, v in kpis.items():
            doc.add_paragraph(f"• {k.title()}: {v}", style='List Bullet')
    else:
        doc.add_paragraph("No KPIs provided")
    
    # ========== KEY FINDINGS ==========
    doc.add_heading("Key Findings & Bottlenecks", level=1)
    
    bottlenecks = final_report.get("consensus_bottlenecks", [])
    if bottlenecks:
        for idx, b in enumerate(bottlenecks, 1):
            doc.add_heading(f"{idx}. {b.get('name', 'Unnamed Bottleneck')}", level=2)
            
            p = doc.add_paragraph()
            p.add_run("Source: ").bold = True
            p.add_run(b.get('source', 'Unknown').replace('_', ' ').title())
            
            p = doc.add_paragraph()
            p.add_run("Priority: ").bold = True
            priority_run = p.add_run(b.get('priority', 'medium').upper())
            if b.get('priority') == 'high':
                priority_run.font.color.rgb = RGBColor(255, 0, 0)
            
            p = doc.add_paragraph()
            p.add_run("Diagnosis: ").bold = True
            p.add_run(b.get('diagnosis', ''))
            
            if b.get("tactical_fix"):
                p = doc.add_paragraph()
                p.add_run("Tactical Fixes:").bold = True
                for fix in b["tactical_fix"]:
                    doc.add_paragraph(f"→ {fix}", style='List Bullet 2')
    else:
        doc.add_paragraph("No bottlenecks identified")
    
    # ========== ACTION PLAN ==========
    doc.add_heading("Recommended Action Plan", level=1)
    
    action_plan = final_report.get("action_plan", [])
    if action_plan:
        for i, a in enumerate(action_plan, start=1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. ").bold = True
            p.add_run(a.get('fix', ''))
            p.add_run(f" (Coach: {a.get('from', '').replace('_', ' ').title()})")
            p.paragraph_format.left_indent = Inches(0.25)
    else:
        doc.add_paragraph("No action plan generated")
    
    # ========== KPIs TO TRACK ==========
    doc.add_heading("Key Performance Indicators (KPIs)", level=1)
    
    doc.add_heading("Essential KPIs to Track", level=2)
    kpis_track = final_report.get("kpis_to_track", [])
    if kpis_track:
        for k in kpis_track:
            doc.add_paragraph(f"• {k.replace('_', ' ').title()}", style='List Bullet')
    else:
        doc.add_paragraph("No KPIs specified")
    
    proposed_kpis = final_report.get("proposed_kpis", [])
    if proposed_kpis:
        doc.add_heading("Additional Proposed KPIs", level=2)
        for p in proposed_kpis:
            para = doc.add_paragraph()
            para.add_run(f"• {p.get('kpi', '').title()}: ").bold = True
            para.add_run(p.get('why', ''))
    
    # ========== COACH INSIGHTS ==========
    doc.add_heading("Detailed Coach Insights", level=1)
    
    coach_insights = final_report.get("coach_insights", {})
    for coach, payload in coach_insights.items():
        coach_name = coach.replace('_', ' ').title()
        doc.add_heading(f"💡 {coach_name}", level=2)
        
        analysis = payload.get("analysis", {})
        if isinstance(analysis, dict):
            # Top Recommendation
            top_rec = analysis.get('top_recommendation', '')
            if top_rec:
                p = doc.add_paragraph()
                p.add_run("Top Recommendation: ").bold = True
                p.add_run(top_rec)
            
            # Summary
            summary = analysis.get('summary', '')
            if summary:
                p = doc.add_paragraph()
                p.add_run("Summary: ").bold = True
                p.add_run(summary)
            
            # Evidence Used
            provenance = payload.get("provenance", [])
            if provenance:
                p = doc.add_paragraph()
                p.add_run(f"Evidence Sources: {len(provenance)} document(s)").italic = True
                p.paragraph_format.left_indent = Inches(0.25)
        else:
            doc.add_paragraph("(Unstructured analysis)")
    
    # ========== FINAL SUMMARY ==========
    doc.add_heading("Consolidated Summary", level=1)
    
    final_summary = final_report.get("final_summary", "")
    if final_summary:
        # Split by || and format each coach's summary
        summaries = final_summary.split(" || ")
        for summary in summaries:
            if summary.strip():
                doc.add_paragraph(f"• {summary.strip()}", style='List Bullet')
    else:
        doc.add_paragraph("No summary available")
    
    # ========== APPENDIX ==========
    doc.add_page_break()
    doc.add_heading("Appendix: RAG Evidence Provenance", level=1)
    
    rag = final_report.get("rag_provenance", {})
    for coach, prov in rag.items():
        coach_name = coach.replace('_', ' ').title()
        doc.add_heading(f"{coach_name} Evidence", level=2)
        
        if prov:
            doc.add_paragraph(f"Total evidence documents retrieved: {len(prov)}")
            for idx, p in enumerate(prov[:5], 1):  # Show top 5
                doc.add_paragraph(
                    f"{idx}. Source: {p.get('source', 'unknown')}, Chunk: {p.get('chunk_id', 'N/A')}, Rank: {p.get('evidence_rank', 'N/A')}"
                )
        else:
            doc.add_paragraph("No RAG evidence retrieved for this coach")
    
    # Save document
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"✅ Saved professional DOCX report: {out}")

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 3:
        print("Usage: python scripts/generate_report_docx.py <final_report.json> <out.docx>")
        raise SystemExit(1)
    fr = json.loads(open(sys.argv[1], encoding="utf-8").read())
    write_consulting_report(fr, sys.argv[2])
